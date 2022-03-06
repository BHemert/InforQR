import os
from turtle import width
import qrcode
import docx
from docx.shared import Inches
from pickle import TRUE

saveQR = TRUE # Set FALSE if you don't want to save qr images

doc = docx.Document()

list = ["AN1 6","AN2 6","AN3 6","AN4 6", # NORTON
 "AT1 9", "AT2 8", "AT3 8", "AT4 9", "AT5 9", "AT6 9", # THINK

  "AEA 8", "AEB 4", "AEC 3", "AED 5", "AEE 8", "AEF 8", "AEG 9", "AEH 8",  #EWENT / EMINENT
   "AEI 8", "AEJ 5", "AEK 5", "AEL 5", "AEM 5", "AEN 5", "AEO 5", "AEP 5", "AEQ 5",  #EWENT / EMINENT
    "AER 5", "AES 5", "AET 5", "AEU 5", "AEV 10", "AEW 4", "AEX 5", "AEY 6", #EWENT / EMINENT

      "AA1 4", "AA2 4", "AA3 4", "AA4 4", "AA5 4", "AA6 4", "AA7 4", "AA8 4", "AA9 2", "AAX 2", "AAY 2", # ACT
       "AU1 5", "AU2 2", "AU3 2", "AU4 5", # TARGUS
        "AL1 8", "AL2 8", "AL3 8", "AL4 8", "AL5 8", # TP-LINK
         "AS1 4", "AS2 7", "AS3 7", "AS4 7", "AS5 8", "AS6 7", "AS7 7", # Speedlink
          "AHI1 1", "AHI2 1" # HI-PLUS
        ]

def docxQR():
    for loc in list:
        loc = loc.split(" ")
        totNum = loc[1]
        for num in range (int(totNum)):

            run = doc.add_paragraph().add_run()
            run.font.name = 'Calibri'
            run.font.size = docx.shared.Pt(36)
            run.bold = True
            run.add_picture("./QRCODES/" + str(loc[0]) + "/" + str(loc[0]) + str(num) + ".png", width=Inches(0.6), height=Inches(0.6))
            run.add_text("  " + str(loc[0]) + str(num))
            try:
                doc.save('./qr.docx')
            except:
                print("Exception: Close docx file and rerun code!")
                exit()

def makeQR():
    for loc in list:
        loc = loc.split(" ")
        totNum = loc[1]
        for num in range (int(totNum) + 1):
            qr = qrcode.QRCode(
                border = 0
            )
            qr.add_data("LK"+ str(loc[0]) + str(num))
            qr.make(fit=True)
            img = qr.make_image()
            if (os.path.exists("./QRCODES/" + str(loc[0]))) and saveQR:
                img.save("./QRCODES/" + str(loc[0]) + "/" + str(loc[0]) + str(num) + ".png")
                print("saved QR: " + str(loc[0]) + str(num) + ".png")
            elif saveQR:
                os.makedirs("./QRCODES/" + str(loc[0]))
                img.save("./QRCODES/" + str(loc[0]) + "/" + str(loc[0]) + str(num) + ".png")
                print("saved QR: " + str(loc[0]) + str(num) + ".png")

makeQR()
docxQR()